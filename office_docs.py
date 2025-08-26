import datetime
import os.path
import sys


from copy import copy
from typing import Optional, List, Any, Dict, Union, Tuple

import PIL
import PIL.Image
import docx
import docx.document
import docx.text
import docx.text.paragraph
import docx.table
import docx.shared
import fitz
import lxml.etree

import openpyxl
import openpyxl.worksheet
import openpyxl.worksheet.worksheet
import openpyxl.styles
import openpyxl.styles.fills
import wx
from docx.oxml.ns import nsdecls
from openpyxl.utils import get_column_letter
import docx2pdf



#noinspection PyProtectedMember
from docx.oxml import CT_TcPr, parse_xml, CT_Drawing  # , CT_Tbl
import docx.text.run

from openpyxl.cell import Cell

import basic
from basic import Logger
import xml.etree.ElementTree



class WordDocument:
    class FloatDrawingProperties(basic.XMLStorable):
        name:str
        filename: str
        pos_x: float # cm
        pos_x_relative: str
        pos_y: float # cm
        pos_y_relative: str
        angle: float # degrees
        def __init__(self):
            basic.XMLStorable.__init__(self)
            self.name = ''
            self.filename = ''
            self.pos_x = 0.0
            self.pos_x_relative = ''
            self.pos_y = 0.0
            self.pos_y_relative = ''
            self.angle = 0.0

        @property
        def filename_only(self)->Optional[str]:
            if self.filename is not None:
                fpath, fname = os.path.split(self.filename)
                return fname
            return None

    class TextFieldInfo:
        name: Optional[str]
        tag: Optional[str]
        value: Optional[str]
        def __init__(self):
            self.name = None
            self.tag = None
            self.value = None

    doc: Optional[docx.document.Document]
    is_opened: bool
    paragraph_tables: Dict[int, List[int]]
    logger: Logger
    _floating_images: Dict[str, Tuple[docx.oxml.CT_Drawing, FloatDrawingProperties]]
    _text_fields: Dict[str, List[xml.etree.ElementTree.Element]]

    def __init__(self, logger: Logger):
        self._floating_images = {}
        self._text_fields = {}

        self.logger = logger
        self.is_opened = False
        self.doc = None
        self.paragraph_tables = {}


    #region Основные функции

    def open(self, file_name: str)->bool:
        if os.path.isfile(file_name):
            try:
                self.doc = docx.Document(file_name)
                self._floating_images.clear()
                self._text_fields.clear()
                self.is_opened = True
                return True
            except Exception as ex:
                self.logger.error(f'Ошибка открытия {file_name} {ex}')

        else:
            self.logger.error(f'Файл {file_name} не найден')
        return False


    def fill_doc_info(self):
        if self.is_opened:
            for drawing in self._get_floating_drawings():
                draw_info = self._get_floating_drawing_properties(drawing)
                if draw_info is not None:
                    self._floating_images[draw_info.name] = drawing, draw_info

            for text_field in self._get_text_fields():
                text_info = self._get_text_field_info(text_field)
                if text_info is not None:
                    if text_info.name not in self._text_fields.keys():
                        self._text_fields[text_info.name] = []
                    self._text_fields[text_info.name].append(text_field)



    def save(self, file_name: str)->bool:
        if self.is_opened:
            try:
                self.doc.save(file_name)
                return True
            except Exception as ex:
                self.logger.error(f'Ошибка сохранения {file_name} {ex}')

        else:
            self.logger.error(f'Нет открытого документа')
        return False

    #endregion


    #region Таблицы
    @property
    def tables(self):
        return self.doc.tables

    def set_cell_value(self, table_index: int, row: int, col: int, value: str):
        if self.is_opened:
            if len(self.doc.tables[table_index].cell(row, col).paragraphs) == 1:
                if len(self.doc.tables[table_index].cell(row, col).paragraphs[0].runs) == 1:
                    self.doc.tables[table_index].cell(row, col).paragraphs[0].runs[0].text = value
                else:
                    self.logger.error(f'Ошибка ячейки таблица={table_index} r={row} c={col} не найден текст')
            else:
                self.logger.error(f'Ошибка ячейки таблица={table_index} r={row} c={col} не найден параграф')

        else:
            self.logger.error(f'Нет открытого документа')

    def clear_cell(self, table_index: int, row: int, col: int):
        c = self.doc.tables[table_index].cell(row, col)
        for par in c.paragraphs:
            self.delete_paragraph(par)
        c.add_paragraph()
        c.paragraphs[0].add_run()

    def set_cell_background_color(self, table_index: int, row: int, col: int, color_value: str):
        cell = self.doc.tables[table_index].cell(row, col)
        # noinspection PyProtectedMember
        tbl_cell = cell._tc
        tbl_cell_properties: CT_TcPr = tbl_cell.get_or_add_tcPr()
        # noinspection PyStringFormat
        shading_elm_1 = parse_xml(f'<w:shd {{}} w:fill="{color_value}"/>'.format(nsdecls('w')))
        tbl_cell_properties.append(shading_elm_1)

    def set_cell_picture(self, table_index: int, row: int, col: int, file_name: str, width: float, height: float):
        if self.is_opened:
            if len(self.doc.tables[table_index].cell(row, col).paragraphs) == 1:
                if len(self.doc.tables[table_index].cell(row, col).paragraphs[0].runs) == 1:

                    width_img = docx.shared.Inches(width)
                    height_igm = docx.shared.Inches(height)
                    self.doc.tables[table_index].cell(row, col).paragraphs[0].runs[0].text = ''
                    self.doc.tables[table_index].cell(row, col).paragraphs[0].runs[0].add_picture(file_name, width_img, height_igm)
                else:
                    self.logger.error(f'Ошибка ячейки {table_index} r={row} c={col}')
            else:
                self.logger.error(f'Ошибка ячейки {table_index} r={row} c={col}')

        else:
            self.logger.error(f'Нет открытого документа')

    def delete_table(self, table_index: int):
        if self.is_opened:
            if 0<=table_index<len(self.doc.tables):
                table = self.doc.tables[table_index]
                # noinspection PyProtectedMember
                table._element.getparent().remove(table._element)
            else:
                self.logger.error(f'Неверный индекс таблицы {table_index}')

        else:
            self.logger.error(f'Документ не открыт')

    #endregion

    #region Параграфы
    @property
    def paragraphs(self):
        return self.doc.paragraphs

    def delete_paragraph(self, paragraph: docx.text.paragraph.Paragraph):
        if self.is_opened:
            paragraph.clear()
            # noinspection PyProtectedMember
            p = paragraph._element
            p.getparent().remove(p)
            # noinspection PyProtectedMember
            paragraph._p = paragraph._element = None
        else:
            self.logger.error(f'Документ не открыт')

    def delete_paragraph_by_index(self, paragraph_index: int):
        if self.is_opened:
            if 0<=paragraph_index<len(self.doc.paragraphs):
                paragraph = self.doc.paragraphs[paragraph_index]
                paragraph.clear()
                #if paragraph in self.doc.paragraphs:
                # noinspection PyProtectedMember
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = paragraph._element = None
                #else:
                #    mlogger.error(f'Параграф {paragraph} не найден в документе {self.doc}')

            else:
                self.logger.error(f'Неверный индекс параграфа {paragraph_index}')

        else:
            self.logger.error(f'Документ не открыт')

    #endregion

    #region Текстовые поля



    def _get_text_fields(self)->List[xml.etree.ElementTree.Element]:
        answer:List[lxml.etree.Element] = []
        if self.is_opened:
            # noinspection PyProtectedMember
            all_fields = self.doc._element.findall('.//{*}sdt')
            for all_field in all_fields:
                answer.append(all_field)
        return answer

    def _get_text_field_info(self, text_field: xml.etree.ElementTree.Element)->Optional[TextFieldInfo]:
        text_field_name = None
        text_field_tag = None
        text_field_value = None
        if self.is_opened:
            # noinspection PyProtectedMember
            root = self.doc._element
            alias_item = text_field.findall('.//{*}sdtPr//{*}alias')
            if alias_item is not None and len(alias_item)==1:
                text_field_name = alias_item[0].get(f'{{{root.nsmap['w']}}}val')
            tag_item = text_field.findall('.//{*}sdtPr//{*}tag')
            if tag_item is not None and len(tag_item)==1:
                text_field_tag = tag_item[0].get(f'{{{root.nsmap['w']}}}val')
            #noinspection PyTypeChecker
            content_items: List[docx.text.run.Run]  = text_field.findall('.//{*}sdtContent//{*}r')
            if content_items is not None and len(content_items) == 1:
                text_field_value = content_items[0].text
            text_info = WordDocument.TextFieldInfo()
            text_info.name = text_field_name
            text_info.tag = text_field_tag
            text_info.value = text_field_value
            return text_info
        return None

    def _set_text_field_value(self, text_field: xml.etree.ElementTree.Element, value: str):
        if self.is_opened:
            # noinspection PyTypeChecker
            content_items: List[docx.text.run.Run]  = text_field.findall('.//{*}sdtContent//{*}r')
            if content_items is not None and len(content_items) == 1:
                content_items[0].text = value
                return True
        return False

    def get_text_field_names(self)->List[str]:
        return list(self._text_fields.keys())

    def set_text_field_value(self, field_name: str, value: str):
        if value is None:
            value = ''
        if field_name in self._text_fields.keys():
            for xml_element in self._text_fields[field_name]:
                self._set_text_field_value(xml_element, value)


    #endregion

    #region Работа с графическими объектами уровня изображения

    @staticmethod
    def _pt_to_cm(val: float):
        return val * 2.54

    @staticmethod
    def _cm_to_pt(val: float):
        return val/2.54



    def _get_floating_drawings(self)->List[CT_Drawing]:
        answer = []
        if self.is_opened:
            # noinspection PyProtectedMember
            drawings: List[docx.oxml.CT_Drawing] = self.doc._element.findall('.//{*}drawing')
            for drawing in drawings:
                props = self._get_floating_drawing_properties(drawing)
                if props:
                    answer.append(drawing)
        return answer

    def _get_floating_drawing_properties(self, drawing: docx.oxml.drawing.CT_Drawing)->Optional[FloatDrawingProperties]:
        image_name = None
        file_name = None
        image_pos_x = None
        image_pos_x_relative = None

        image_pos_y = None
        image_pos_y_relative = None
        image_angle = None
        if self.is_opened:
            doc_prs: docx.oxml.shape.CT_NonVisualDrawingProps = drawing.findall('.//{*}anchor//{*}docPr')
            if doc_prs is not None and len(doc_prs)==1:
                try:
                    if 'name' in doc_prs[0].attrib:
                        image_name = doc_prs[0].attrib['name']
                    if 'descr' in doc_prs[0].attrib:
                        file_name  =doc_prs[0].attrib['descr']

                    pos_offset = drawing.findall('.//{*}anchor//{*}positionH//{*}posOffset')
                    if pos_offset and len(pos_offset) == 1:
                        if 'relativeFrom' in pos_offset[0].attrib:
                            image_pos_x_relative = pos_offset[0].attrib['relativeFrom']

                        image_pos_x = int(pos_offset[0].text)
                        image_pos_x = self._float_image_val_to_cm(image_pos_x)

                    pos_offset = drawing.findall('.//{*}anchor//{*}positionV//{*}posOffset')
                    if pos_offset and len(pos_offset) == 1:
                        if 'relativeFrom' in pos_offset[0].attrib:
                            image_pos_y_relative = pos_offset[0].attrib['relativeFrom']

                        image_pos_y = int(pos_offset[0].text)
                        image_pos_y = self._float_image_val_to_cm(image_pos_y)

                    angle_offset = drawing.findall('.//{*}anchor//{*}graphic//{*}graphicData//{*}pic//{*}spPr//{*}xfrm')
                    if angle_offset and len(angle_offset) == 1:
                        if 'rot' in angle_offset[0].attrib:
                            image_angle = int(angle_offset[0].attrib['rot'])
                            image_angle = self._float_image_val_to_angle(image_angle)
                    img_info = WordDocument.FloatDrawingProperties()
                    img_info.name = image_name
                    img_info.filename = file_name
                    img_info.pos_x = image_pos_x
                    img_info.pos_x_relative = image_pos_x_relative
                    img_info.pos_y = image_pos_y
                    img_info.pos_y_relative = image_pos_y_relative
                    img_info.angle = image_angle
                    return img_info
                except Exception as ex:
                    self.logger.error(f'Ошибка чтения данных изображения {ex}')
        return None

    def _set_floating_drawing_properties(self, drawing: docx.oxml.drawing.CT_Drawing, props: FloatDrawingProperties):
        doc_prs: docx.oxml.shape.CT_NonVisualDrawingProps = drawing.findall('.//{*}anchor//{*}docPr')
        result = True
        if doc_prs is not None and len(doc_prs) == 1:
            try:
                pos_offset = drawing.findall('.//{*}anchor//{*}positionH//{*}posOffset')
                if pos_offset and len(pos_offset) == 1:
                    pos_offset[0].text = str(self._float_image_val_from_cm(props.pos_x))

                pos_offset = drawing.findall('.//{*}anchor//{*}positionV//{*}posOffset')
                if pos_offset and len(pos_offset) == 1:
                    pos_offset[0].text = str(self._float_image_val_from_cm(props.pos_y))

                angle_offset = drawing.findall('.//{*}anchor//{*}graphic//{*}graphicData//{*}pic//{*}spPr//{*}xfrm')
                if angle_offset and len(angle_offset) == 1:
                    angle_offset[0].attrib['rot'] = str(self._float_image_val_from_angle(props.angle))
            except Exception as ex:
                self.logger.error(f'Ошибка чтения данных изображения {ex}')
        return result

    def is_floating_image_exists(self, image_name:str):
        return image_name in self._floating_images.keys()
    def get_floating_drawing_info(self)->List[FloatDrawingProperties]:
        answ = []
        for img_name, data in self._floating_images.items():
            answ.append(data[1])
        return answ

    def set_floating_drawing_properties(self, image_name: str, props: FloatDrawingProperties):
        if image_name in self._floating_images.keys():
            self._set_floating_drawing_properties(self._floating_images[image_name][0], props)
    @staticmethod
    def _float_image_val_from_cm(val: float):
        return int(val * 360045)

    @staticmethod
    def _float_image_val_to_cm(val: float):
        return val / 360045

    @staticmethod
    def _float_image_val_to_angle(val: float):
        return val / 60000

    @staticmethod
    def _float_image_val_from_angle(val: float):
        return int(val * 60000)

    #endregion

    class Writer(object):
        log = []

        def write(self, data):
            pass

    def convert_to_pdf(self, filename: str, outputfilename: str):
        if os.path.exists(filename):
            if not os.path.exists(outputfilename):
                try:
                    old_output = sys.stderr
                    sys.stderr = self.Writer()
                    docx2pdf.convert(filename,outputfilename)
                    sys.stderr = old_output
                    if os.path.exists(outputfilename):
                        return True
                    return False
                except Exception as ex:
                    self.logger.error(f'Ошибка преобразования {filename} {ex}')

            else:
                self.logger.error(f'Файл {outputfilename} уже существует')
                return False
        else:
            self.logger.error(f'Файл {filename} не найден')
            return False


    #region Дополнительные функции
    @staticmethod
    def _iter_block_items(parent):
        # noinspection PyProtectedMember
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
            # noinspection PyProtectedMember
        elif isinstance(parent, docx.table._Cell):
            # noinspection PyProtectedMember
            parent_elm = parent._tc
        elif isinstance(parent, docx.table._Row):
            # noinspection PyProtectedMember
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, docx.oxml.CT_P):
                yield docx.text.paragraph.Paragraph(child, parent)
            elif isinstance(child, docx.oxml.CT_Tbl):
                yield docx.table.Table(child, parent)

    def analyse_structure(self):
        self.paragraph_tables = {}
        table_index = -1
        paragraph_index = -1
        if self.is_opened:
            for block_item in self._iter_block_items(self.doc):
                if isinstance(block_item, docx.table.Table):
                    #table = block_item
                    table_index +=1
                    if paragraph_index not in self.paragraph_tables.keys():
                        self.paragraph_tables[paragraph_index] = []
                    self.paragraph_tables[paragraph_index].append(table_index)
                elif isinstance(block_item, docx.text.paragraph.Paragraph):
                    #paragraph: docx.text.paragraph.Paragraph = block_item
                    paragraph_index += 1
                else:
                    self.logger.error(f'Неизвестный тип блока {block_item} в параграфе {paragraph_index}')

        else:
            self.logger.error(f'Нет открытого документа')

    @staticmethod
    def get_excel_colour(color: wx.Colour) -> str:
        output_str = f'{color.GetRed():0{2}x}{color.GetGreen():0{2}x}{color.GetBlue():0{2}x}'.upper()
        return output_str

    #endregion

class ExcelDocument:
    xls_name: Optional[str]
    workbook: Optional[openpyxl.Workbook]
    worksheet: Optional[openpyxl.worksheet.worksheet.Worksheet]
    is_opened: bool
    logger: Logger

    def __init__(self, logger: Logger):
        self.logger = logger
        self.is_opened = False
        self.doc = None
        self.xls_name = ''

    def create(self):
        """ Создать новый файл """
        self.worksheet = None
        self.is_opened = False
        try:
            self.workbook = openpyxl.Workbook()
            self.worksheet = self.workbook.active
            self.is_opened = True
        except Exception as ex:
            self.logger.error(f'Ошибка создания xlsx файла {ex}')

    def open(self, file_name: str):
        """ Прочитать книгу по имени файла """
        self.is_opened = False
        self.worksheet = None
        self.xls_name = os.path.split(file_name)[1]
        if os.path.isfile(file_name):
            try:
                self.workbook = openpyxl.load_workbook(file_name)
                self.worksheet = self.workbook.active
                self.is_opened = True
            except Exception as ex:
                self.logger.error(f'Ошибка загрузки файла {self.xls_name}: {ex}')
                self.xls_name = None
        else:
            self.logger.error(f'Файл {file_name} не найден')
        return self.is_opened

    def save(self, file_name: str):
        """ Сохранить текущий файл под заданным именем файла """
        if self.is_opened:
            try:
                self.workbook.save(file_name)
            except Exception as ex:
                self.logger.error(f'Ошибка сохранения файла {os.path.split(file_name)[1]}: {ex}')
        else:
            self.logger.error(f'Нет открытого документа')

    def close(self):
        """ Закрыть текущий файл """
        if self.is_opened:
            self.workbook.close()
        else:
            self.logger.error(f'Нет открытого документа')

    def open_worksheet(self, worksheet_name: str):
        """ Перейти/открыть лист с заданным именем """
        if self.is_opened:
            sheets: List[str] = self.workbook.sheetnames
            if worksheet_name in sheets:
                self.worksheet = self.workbook[sheets[sheets.index(worksheet_name)]]
                self.workbook.active = self.worksheet
                return True
            else:
                self.logger.error(f'Лист {worksheet_name} не найден')
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    def create_worksheet(self, worksheet_name:str):
        if self.is_opened:
            if not self.is_worksheet_exits(worksheet_name):
                self.workbook.create_sheet(worksheet_name)
                return True
            else:
                self.logger.error(f'Лист {worksheet_name} уже существует')
        else:
            self.logger.error('Нет открытого документа')
        return False

    def is_worksheet_exits(self, worksheet_name: str):
        """ Проверить существует ли лист """
        if self.is_opened:
            return worksheet_name in self.workbook.sheetnames
        return False

    def delete_worksheet(self, worksheet_name: str):
        """ удалить лист под заданным имененм """
        if self.is_opened:
            sheets: List[str] = self.workbook.sheetnames
            if worksheet_name in sheets:
                del_sheet = self.workbook.get_sheet_by_name(worksheet_name)
                self.workbook.remove_sheet(del_sheet)
                return True
            else:
                self.logger.error(f'Лист {worksheet_name} не найден')
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    def rename_worksheet(self, old_name: str, new_name: str):
        """ Переименовать лист из старого названия в новое """
        if self.is_worksheet_exits(old_name):
            ss_sheet = self.workbook.get_sheet_by_name(old_name)
            ss_sheet.title = new_name
        else:
            self.logger.error(f'Лист {old_name} не найден')

    def delete_row(self, row_index: int):
        """ Удалить заданную строку """
        if self.is_opened:
            self.worksheet.delete_rows(row_index, 1)

    def get_sheet_names(self):
        """ Получить список листов """
        if self.is_opened:
            return self.workbook.get_sheet_names()
        return []

    def get_data(self, row: int, col: int)->Optional[Union[str, datetime.datetime, float, int, bool]]:
        """ Считать значение из ячейки """
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            return cell.value
        else:
            self.logger.error(f'Нет открытого документа')
        return None

    def set_border(self, row: int, col: int, value: openpyxl.styles.borders.Border):
        """ Установить границы для ячейки """
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            cell.border = value
            return True
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    def add_border(self, row: int, col:int, value: openpyxl.styles.Border):
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            old_border_l = cell.border.left
            old_border_r = cell.border.right
            old_border_t = cell.border.top
            old_border_b = cell.border.bottom
            s: openpyxl.styles.Side
            if value.left.style is not None:
                old_border_l = value.left
            if value.right.style is not None:
                old_border_r = value.right
            if value.top.style is not None:
                old_border_t = value.top
            if value.bottom.style is not None:
                old_border_b = value.bottom
            cell.border = openpyxl.styles.Border(left=old_border_l,right=old_border_r, top=old_border_t, bottom=old_border_b)
            return True
        else:
            self.logger.error(f'Нет открытого документа')
        return False


    def set_col_width(self, col: int, width: int):
        self.worksheet.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width

    def set_row_height(self, row:int, height: int):
        self.worksheet.row_dimensions[row].height = height

    def set_data(self, row: int, col: int, value: Any, cell_format: str='@', quote_prefix: bool = False):
        """ Установить значение в ячейку """
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)

            if cell_format == 'auto':
                if type(value) == datetime.date:
                    cell_format = 'dd.mm.yyyyy'
                elif type(value)==int:
                    cell_format = '0'
                elif type(value)==float:
                    cell_format = '0,0000'
                elif type(value)==str:
                    cell_format = '@'
            if quote_prefix:
                normal_style = None
                if 'Normal' in self.workbook.named_styles:
                    normal_style = 'Normal'
                else:
                    normal_style = openpyxl.styles.NamedStyle(name="Normal", number_format='@')
                cell.quotePrefix = True
                cell.style = normal_style
            cell.number_format = cell_format
            cell.value = value
            return True
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    def get_fill(self, row: int, col: int)->Optional[openpyxl.styles.fills.PatternFill]:
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            return cell.fill
        else:
            self.logger.error(f'Нет открытого документа')
        return None

    def get_row_count(self, sheet_name: str):
        if self.is_opened:
            current_sheet = self.workbook.active
            if self.open_worksheet(sheet_name):
                answer = self.worksheet.max_row
                self.open_worksheet(current_sheet)
                return answer
        return -1

    def set_fill(self, row: int, col: int, fill: openpyxl.styles.fills.Fill):
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            cell.fill = copy(fill)
            return True
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    def set_alignment(self, row: int, col: int, alignment: openpyxl.styles.Alignment):
        if self.is_opened:
            cell: openpyxl.cell.Cell = self.worksheet.cell(row, col)
            cell.alignment = copy(alignment)
            return True
        else:
            self.logger.error(f'Нет открытого документа')
        return False

    @staticmethod

    def copy_row(src_xls: 'ExcelDocument', src_row: int, dst_xls: 'ExcelDocument', dst_row:int, col: int, value: bool=True, font: bool = True, border:bool = True, fill:bool = True, number_format: bool = True, data_type:bool = True, protection: bool = False, alignment: bool = True, width: bool = True):
        if width:
            dst_xls.worksheet.column_dimensions[get_column_letter(col)].width = src_xls.worksheet.column_dimensions[get_column_letter(col)].width

        if value:
            dst_xls.worksheet.cell(dst_row, col).value = copy(src_xls.worksheet.cell(src_row, col).value)
        if data_type:
            dst_xls.worksheet.cell(dst_row, col).data_type = copy(src_xls.worksheet.cell(src_row, col).data_type)
        if font:
            src_font: openpyxl.styles.Font = src_xls.worksheet.cell(src_row, col).font
            new_font = openpyxl.styles.Font(name=src_font.name,
                                            size=src_font.size,
                                            bold=src_font.bold,
                                            italic=src_font.italic,
                                            vertAlign=src_font.vertAlign,
                                            underline=src_font.underline,
                                            strike=src_font.strike,
                                            color=src_font.color)
            dst_xls.worksheet.cell(dst_row, col).font = new_font
        if border:
            dst_xls.worksheet.cell(dst_row, col).border = copy(src_xls.worksheet.cell(src_row, col).border)
        if fill:
            dst_xls.worksheet.cell(dst_row, col).fill = copy(src_xls.worksheet.cell(src_row, col).fill)
        if number_format:
            dst_xls.worksheet.cell(dst_row, col).number_format = copy(src_xls.worksheet.cell(src_row, col).number_format)
        if protection:
            dst_xls.worksheet.cell(dst_row, col).protection = copy(src_xls.worksheet.cell(src_row, col).protection)
        if alignment:
            dst_xls.worksheet.cell(dst_row, col).alignment = copy(src_xls.worksheet.cell(src_row, col).alignment)

    def get_cell_fill(self, color: wx.Colour, style: int)->openpyxl.styles.PatternFill:
        if style in [wx.BRUSHSTYLE_SOLID]:
            color_str = self.convert_color(color)
            return openpyxl.styles.PatternFill(start_color=color_str, end_color=color_str, fill_type='solid')
        else:
            self.logger.error(f'{self} неизвестный стиль {style}')

    def get_cell_borders(self, top: str = 'thin', bottom:str='thin', left:str='thin', right: str='thin'):
        kwargs = {}
        if left:
            kwargs['left'] = openpyxl.styles.borders.Side(style=left)
        if right:
            kwargs['right'] = openpyxl.styles.borders.Side(style=right)
        if top:
            kwargs['top'] = openpyxl.styles.borders.Side(style=top)
        if bottom:
            kwargs['bottom'] = openpyxl.styles.borders.Side(style=bottom)
        return openpyxl.styles.borders.Border(**kwargs)

    def get_cell_alignment(self, horizontal:str='center', vertical:str='center', wrap_text: bool=True):
        #horizontal = left, right, center
        #vertical = top, bottom, center
        kwargs = {}
        if horizontal:
            kwargs['horizontal'] = horizontal
        if vertical:
            kwargs['vertical'] = vertical
        if wrap_text:
            kwargs['wrap_text'] = True
        return openpyxl.styles.Alignment(**kwargs)

    def merge_cell(self, row: int, col: int, width:int, height: int ):
        self.worksheet.merge_cells(start_row=row, start_column=col, end_row=row+height, end_column=col+width)

    def convert_color(self, color: wx.Colour):
        if color:
            return f'00{color.GetRed():02x}{color.GetGreen():02x}{color.GetBlue():02x}'
        else:
            return '00000000'

class PDF:

    @staticmethod
    def pdf_rasterize(filename:str, resolution: int = 100):
        pdf_doc:fitz.Document = fitz.open(filename)
        images: List[PIL.Image] = []
        # noinspection PyTypeChecker
        for i, p in enumerate(pdf_doc):
            page = pdf_doc.load_page(i)
            pix = page.get_pixmap(dpi=resolution)
            img = PIL.Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            images.append(img)
        images[0].save(filename, "PDF", resolution=resolution, save_all=True, append_images=images[1:])

    def combine_pdf(self, src_filename: str, dst_filenames: List[str]):
        pass








